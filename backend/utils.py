# -*- coding:utf-8 -*-
# --------------------------------------------------------
# Copyright (C), 2016-2020, omosoft, All rights reserved
# --------------------------------------------------------
# @Name:        utils.py
# @Author:      lizhe
# @Created:     2025/3/16 - 16:35
# --------------------------------------------------------
import dataclasses
import os
from typing import Dict, Optional

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from loguru import logger
from werkzeug.datastructures import FileStorage

from structs import LanguageType
from translate import deepseek_translate, deepl_translate

ERROR_CODE = 40000
ERROR_RESPONSE_CODE = 400
MAX_SIZE = 10 * 1024 * 1024  # 100MB
ALLOWED_EXTENSIONS = {'.docx', '.txt', '.md'}


@dataclasses.dataclass
class ResponseObject(object):
    message: str = ""
    link: str = ""
    status: int = 20000


def check_file(file_key: str, files: Dict) -> Optional[Dict]:
    if file_key not in files:
        response_object = ResponseObject()
        response_object.message = f"没有选择{file_key}文件"
        response_object.status = ERROR_CODE
        return dataclasses.asdict(response_object)  # type: ignore


def check_file_available(file: FileStorage, ext: str) -> Optional[Dict]:
    if file.filename == '':
        response_object = ResponseObject()
        response_object.message = f"没有选择文件"
        response_object.status = ERROR_CODE
        return dataclasses.asdict(response_object)  # type: ignore
    if file.content_length > MAX_SIZE:
        response_object = ResponseObject()
        response_object.message = f"文件大小超过限制"
        response_object.status = ERROR_CODE
        return dataclasses.asdict(response_object)  # type: ignore
    if ext.lower() not in ALLOWED_EXTENSIONS:
        response_object = ResponseObject()
        response_object.message = f"仅支持{ALLOWED_EXTENSIONS}类型，当前类型是{ext}"
        response_object.status = ERROR_CODE
        return dataclasses.asdict(response_object)  # type: ignore


def read_content(abs_file: str, ext: str) -> str:
    if ext == ".docx":
        contents = read_doc_content(abs_file)
    else:
        with open(abs_file, "r", encoding="utf-8") as f:
            contents = f.read()
    return contents


def translate_content(content: str, model: str, language: str) -> str:
    logger.debug(f"use {model} and translate to {language}")
    if model.lower() == "deepseek":
        content = deepseek_translate(content, language)
    elif model.lower() == "deepl":
        language_type = LanguageType.from_name(language)
        content = deepl_translate(content, language_type.value[-1])
    else:
        content = f"{model} not support"
    return content


def get_output_file(name: str):
    folder_name = os.path.join(os.getcwd(), "temp")
    os.makedirs(folder_name, exist_ok=True)
    template_file = os.path.join(folder_name, name)
    logger.debug(f"{template_file = }")
    return template_file


def read_doc_content(file_path: str) -> str:
    """
    读取 .docx 文件的内容。

    参数:
    file_path (str): 文件的绝对路径。

    返回:
    List[str]: 文件的所有内容。
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件 {file_path} 不存在。")

    if file_path.endswith('.docx'):
        doc = Document(file_path)
        content = [para.text for para in doc.paragraphs]
    else:
        raise ValueError("仅支持 .docx 文件格式。")

    return "\n".join(content)


def write_doc_content(contents: str, file_name: str):
    """
    将内容写入到 .docx 文件中。

    参数:
    contents (List[str]): 要写入的内容列表。
    file_name (str): 输出文件的名称。
    """
    if not file_name.endswith('.docx'):
        raise RuntimeError("file type not valid")
    doc = Document()
    doc.add_paragraph(contents)
    doc.save(file_name)


def write_markdown_content(contents: str, file_name: str):
    """
    写入markdown形式的文件
    :param contents: 要写入的内容列表。
    :param file_name: 输出文件的名称。
    :return:
    """
    html = markdown.markdown(contents)
    soup = BeautifulSoup(html, 'html.parser')
    # 创建Word文档
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 处理HTML元素
    for element in soup:
        if element.name == 'h1':
            p = doc.add_heading('', level=1)
            p.add_run(element.text)
        elif element.name == 'h2':
            p = doc.add_heading('', level=2)
            p.add_run(element.text)
        elif element.name == 'h3':
            p = doc.add_heading('', level=3)
            p.add_run(element.text)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
        elif element.name == 'blockquote':
            p = doc.add_paragraph()
            p.add_run(element.text).italic = True
        elif element.name == 'pre':
            # 处理代码块
            code = element.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Pt(20)

    # 保存Word文档
    doc.save(file_name)
